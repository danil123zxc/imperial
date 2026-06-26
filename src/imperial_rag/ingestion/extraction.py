from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol

import fitz
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document
from openpyxl import load_workbook
from striprtf.striprtf import rtf_to_text

from imperial_rag.ingestion.manifest import FileRecord, FileStatus
from imperial_rag.ingestion.ocr import OcrCache, OcrResult


ARCHIVE_EXTENSIONS = {".rar", ".zip", ".7z"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff"}


class SupportsOcr(Protocol):
    def extract_image_text(self, image_path: Path) -> OcrResult:
        ...


@dataclass(frozen=True)
class ExtractionResult:
    record: FileRecord
    status: FileStatus
    documents: list[Document]
    extraction_method: str | None = None
    message: str = ""


def _base_metadata(record: FileRecord, source_type: str) -> dict[str, str | int | None]:
    return {
        "file_id": record.file_id,
        "file_path": str(record.absolute_path),
        "relative_path": str(record.relative_path),
        "file_name": record.filename,
        "file_extension": record.extension,
        "file_hash": record.sha256,
        "duplicate_group_id": record.duplicate_group_id,
        "parent_folder": str(record.parent_folder),
        "inferred_category": record.inferred_category,
        "source_type": source_type,
    }


def _artifact_dir(record: FileRecord, artifact_root: Path | None) -> Path:
    root = artifact_root or record.absolute_path.parent / ".imperial_rag_artifacts"
    target = root / record.sha256
    target.mkdir(parents=True, exist_ok=True)
    return target


def _ocr_image(
    record: FileRecord,
    image_path: Path,
    source_type: str,
    ocr_client: SupportsOcr | None,
    metadata: dict[str, str | int | None],
    ocr_cache: OcrCache | None = None,
    image_id: str = "root",
    warnings: list[str] | None = None,
    empty_label: str | None = None,
) -> list[Document]:
    ocr_result = ocr_cache.lookup(record.sha256, image_id) if ocr_cache is not None else None
    if ocr_result is None:
        if ocr_client is None:
            if warnings is not None and empty_label:
                warnings.append(f"OCR client not configured for {empty_label}")
            return []
        ocr_result = ocr_client.extract_image_text(image_path)
        if ocr_cache is not None and ocr_result.text:
            ocr_cache.store(record.sha256, image_id, ocr_result)
    if not ocr_result.text:
        if warnings is not None and empty_label:
            warnings.append(f"OCR returned empty text for {empty_label}")
        return []
    merged_metadata = _base_metadata(record, source_type)
    merged_metadata.update(metadata)
    merged_metadata["ocr_method"] = ocr_result.method
    merged_metadata["ocr_cached"] = ocr_result.cached
    return [Document(page_content=ocr_result.text, metadata=merged_metadata)]


def _docx_blocks(docx) -> list[Paragraph | Table]:
    blocks: list[Paragraph | Table] = []
    for child in docx.element.body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(Paragraph(child, docx))
        elif isinstance(child, CT_Tbl):
            blocks.append(Table(child, docx))
    return blocks


def _is_heading(paragraph: Paragraph) -> bool:
    style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
    return style_name.casefold().startswith("heading")


def _locator_slug(text: str) -> str:
    slug = re.sub(r"[^\w]+", "-", text.casefold(), flags=re.UNICODE).strip("-")
    return slug or "section"


def _element_hash(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8")).hexdigest()


def _section_parent_id(record: FileRecord, section_index: int, section_slug: str) -> str:
    return f"{record.file_id}:section:{section_index}:{section_slug}"


def _structured_metadata(
    record: FileRecord,
    source_type: str,
    *,
    source_locator: str,
    page_content: str,
    parent_id: str,
    section_heading: str | None = None,
    **extra: str | int | None,
) -> dict[str, str | int | None]:
    metadata = _base_metadata(record, source_type)
    metadata["source_locator"] = source_locator
    metadata["source_doc_id"] = f"{record.file_id}:{source_type}:{source_locator}"
    metadata["parent_id"] = parent_id
    metadata["element_hash"] = _element_hash(page_content)
    metadata["element_id"] = f"{metadata['source_doc_id']}:{str(metadata['element_hash'])[:12]}"
    if section_heading:
        metadata["section_heading"] = section_heading
    metadata.update(extra)
    return metadata


def _table_lines(table: Table) -> list[str]:
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    return lines


def _extract_docx(
    record: FileRecord,
    ocr_client: SupportsOcr | None = None,
    ocr_cache: OcrCache | None = None,
    artifact_root: Path | None = None,
    warnings: list[str] | None = None,
) -> list[Document]:
    docx = DocxDocument(str(record.absolute_path))
    documents: list[Document] = []
    body_lines: list[str] = []
    body_index = 0
    table_index = 0
    section_heading: str | None = None
    section_index = 0
    section_slug = ""

    def flush_body() -> None:
        nonlocal body_lines, body_index
        body_text = "\n".join(line for line in body_lines if line.strip()).strip()
        body_lines = []
        if not body_text:
            return
        body_index += 1
        if section_heading:
            source_locator = f"section:{section_index}:{section_slug}"
            parent_id = _section_parent_id(record, section_index, section_slug)
        else:
            source_locator = f"body:{body_index}"
            parent_id = record.file_id
        documents.append(
            Document(
                page_content=body_text,
                metadata=_structured_metadata(
                    record,
                    "body",
                    source_locator=source_locator,
                    page_content=body_text,
                    parent_id=parent_id,
                    section_heading=section_heading,
                    element_index=body_index,
                ),
            )
        )

    for block in _docx_blocks(docx):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if _is_heading(block):
                flush_body()
                section_heading = text
                section_index += 1
                section_slug = _locator_slug(text)
                continue
            if section_heading and not body_lines:
                body_lines.append(section_heading)
            body_lines.append(text)
            continue

        flush_body()
        lines = _table_lines(block)
        if not lines:
            continue
        table_index += 1
        table_text = "\n".join(lines)
        row_start = 1
        row_end = len(lines)
        if section_heading:
            source_locator = f"section:{section_index}:{section_slug}:table:{table_index}:rows:{row_start}-{row_end}"
            parent_id = _section_parent_id(record, section_index, section_slug)
        else:
            source_locator = f"table:{table_index}:rows:{row_start}-{row_end}"
            parent_id = record.file_id
        documents.append(
            Document(
                page_content=table_text,
                metadata=_structured_metadata(
                    record,
                    "table",
                    source_locator=source_locator,
                    page_content=table_text,
                    parent_id=parent_id,
                    section_heading=section_heading,
                    table_index=table_index,
                    row_start=row_start,
                    row_end=row_end,
                ),
            )
        )
    flush_body()

    if ocr_client is not None or ocr_cache is not None:
        target_dir = _artifact_dir(record, artifact_root)
        with zipfile.ZipFile(record.absolute_path) as archive:
            media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
            for index, media_name in enumerate(media_names, start=1):
                try:
                    suffix = Path(media_name).suffix.lower() or ".img"
                    image_path = target_dir / f"embedded-{index}{suffix}"
                    image_path.write_bytes(archive.read(media_name))
                    documents.extend(
                        _ocr_image(
                            record,
                            image_path,
                            "embedded_image",
                            ocr_client,
                            {"image_index": index, "embedded_media_name": media_name},
                            ocr_cache=ocr_cache,
                            image_id=f"embedded-{index}",
                            warnings=warnings,
                            empty_label=f"embedded image {index}",
                        )
                    )
                except Exception as exc:
                    if warnings is not None:
                        warnings.append(f"embedded image OCR failed for {media_name}: {exc}")
                    continue
    return documents


def _extract_pdf(
    record: FileRecord,
    ocr_client: SupportsOcr | None,
    ocr_cache: OcrCache | None,
    artifact_root: Path | None,
    warnings: list[str] | None = None,
) -> list[Document]:
    documents: list[Document] = []
    target_dir = _artifact_dir(record, artifact_root)
    with fitz.open(record.absolute_path) as pdf:
        for zero_based_page_index in range(pdf.page_count):
            page_index = zero_based_page_index + 1
            page = pdf.load_page(zero_based_page_index)
            raw_text = page.get_text("text")
            text = raw_text.strip() if isinstance(raw_text, str) else str(raw_text).strip()
            if text:
                metadata = _base_metadata(record, "pdf_page")
                metadata["page_number"] = page_index
                documents.append(Document(page_content=text, metadata=metadata))
                continue
            image_path = target_dir / f"{record.absolute_path.stem}-page-{page_index}.jpg"
            pixmap = page.get_pixmap(dpi=200)
            pixmap.save(image_path)
            documents.extend(
                _ocr_image(
                    record,
                    image_path,
                    "pdf_page",
                    ocr_client,
                    {"page_number": page_index, "render_dpi": 200},
                    ocr_cache=ocr_cache,
                    image_id=f"page-{page_index}",
                    warnings=warnings,
                    empty_label=f"page {page_index}",
                )
            )
    return documents


def _extract_xlsx(record: FileRecord) -> list[Document]:
    workbook = load_workbook(record.absolute_path, data_only=True, read_only=True)
    documents: list[Document] = []
    for sheet in workbook.worksheets:
        lines: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                lines.append(" | ".join(cells))
        if lines:
            metadata = _base_metadata(record, "sheet")
            metadata["sheet_name"] = sheet.title
            documents.append(Document(page_content="\n".join(lines), metadata=metadata))
    return documents


def _extract_rtf(record: FileRecord) -> list[Document]:
    text = rtf_to_text(record.absolute_path.read_text(encoding="utf-8", errors="ignore")).strip()
    if not text:
        return []
    return [Document(page_content=text, metadata=_base_metadata(record, "body"))]


def extract_file(
    record: FileRecord,
    ocr_client: SupportsOcr | None = None,
    ocr_cache: OcrCache | None = None,
    artifact_root: Path | None = None,
) -> ExtractionResult:
    if record.extension in ARCHIVE_EXTENSIONS:
        return ExtractionResult(
            record=record,
            status=FileStatus.MANIFEST_ONLY,
            documents=[],
            extraction_method=None,
            message="archive files recorded but not extracted in v1",
        )
    if record.extension == ".docx":
        warnings: list[str] = []
        documents = _extract_docx(
            record,
            ocr_client=ocr_client,
            ocr_cache=ocr_cache,
            artifact_root=artifact_root,
            warnings=warnings,
        )
        return ExtractionResult(
            record,
            FileStatus.INDEXED if documents else FileStatus.NO_TEXT,
            documents,
            "python_docx",
            "; ".join(warnings),
        )
    if record.extension == ".pdf":
        warnings: list[str] = []
        documents = _extract_pdf(
            record,
            ocr_client=ocr_client,
            ocr_cache=ocr_cache,
            artifact_root=artifact_root,
            warnings=warnings,
        )
        return ExtractionResult(
            record,
            FileStatus.INDEXED if documents else FileStatus.NO_TEXT,
            documents,
            "pymupdf",
            "; ".join(warnings),
        )
    if record.extension in IMAGE_EXTENSIONS:
        warnings: list[str] = []
        documents = _ocr_image(
            record,
            record.absolute_path,
            "image",
            ocr_client,
            {"image_hash": record.sha256},
            ocr_cache=ocr_cache,
            image_id="image",
            warnings=warnings,
            empty_label="image",
        )
        return ExtractionResult(
            record,
            FileStatus.INDEXED if documents else FileStatus.NO_TEXT,
            documents,
            "image_ocr",
            "; ".join(warnings),
        )
    if record.extension == ".xlsx":
        documents = _extract_xlsx(record)
        return ExtractionResult(record, FileStatus.INDEXED if documents else FileStatus.NO_TEXT, documents, "openpyxl")
    if record.extension == ".rtf":
        documents = _extract_rtf(record)
        return ExtractionResult(record, FileStatus.INDEXED if documents else FileStatus.NO_TEXT, documents, "striprtf")
    if record.extension == ".doc":
        return ExtractionResult(
            record=record,
            status=FileStatus.UNSUPPORTED,
            documents=[],
            extraction_method=None,
            message="legacy .doc requires a safe local converter; recorded but not extracted in v1",
        )
    return ExtractionResult(
        record=record,
        status=FileStatus.UNSUPPORTED,
        documents=[],
        extraction_method=None,
        message=f"unsupported extension: {record.extension}",
    )
