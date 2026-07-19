import sqlite3
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from imperial_rag.ingestion.extraction import extract_file
from imperial_rag.ingestion.manifest import FileStatus, scan_files
from imperial_rag.ingestion import ocr as ocr_module
from imperial_rag.ingestion.ocr import OcrCache, OcrResult
from openpyxl import Workbook
from PIL import Image


class TrackingConnection:
    def __init__(self, connection: sqlite3.Connection) -> None:
        object.__setattr__(self, "_connection", connection)
        object.__setattr__(self, "closed", False)

    def __getattr__(self, name: str):
        return getattr(self._connection, name)

    def close(self) -> None:
        object.__setattr__(self, "closed", True)
        self._connection.close()


class FakeOcrClient:
    def __init__(self, prefix: str = "OCR") -> None:
        self.prefix = prefix
        self.calls: list[Path] = []

    def extract_image_text(self, image_path: Path) -> OcrResult:
        self.calls.append(image_path)
        return OcrResult(text=f"{self.prefix}:{image_path.name}", method="fake_ocr")


class EmptyOcrClient:
    def extract_image_text(self, image_path: Path) -> OcrResult:
        return OcrResult(text="", method="fake_ocr")


class CannedDashScopeOcrClient:
    def extract_image_text(self, image_path: Path) -> OcrResult:
        return OcrResult(text="No visible text in this image.", method="dashscope:qwen-vl-ocr-test")


def _record_for(path: Path):
    for record in scan_files(path.parent):
        if record.absolute_path == path.resolve():
            return record
    raise AssertionError(f"no scanned record for {path}")


def _make_image(path: Path) -> None:
    Image.new("RGB", (20, 20), "white").save(path)


def test_archive_is_manifest_only(tmp_path):
    path = tmp_path / "archive.rar"
    path.write_bytes(b"archive")
    record = _record_for(path)

    result = extract_file(record)

    assert result.status == FileStatus.MANIFEST_ONLY
    assert result.documents == []
    assert result.extraction_method is None
    assert "archive files recorded but not extracted" in result.message


def test_docx_text_and_table_extract_to_langchain_documents_with_citation_metadata(tmp_path):
    path = tmp_path / "policy.docx"
    docx = DocxDocument()
    docx.add_paragraph("Регламент возврата товара")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Ответственный"
    table.cell(0, 1).text = "Склад"
    docx.save(path)
    record = _record_for(path)

    result = extract_file(record)

    assert result.status == FileStatus.INDEXED
    assert result.extraction_method == "python_docx"
    assert [doc.metadata["source_type"] for doc in result.documents] == ["body", "table"]
    assert "Регламент возврата товара" in result.documents[0].page_content
    assert "Ответственный | Склад" in result.documents[1].page_content
    for document in result.documents:
        assert document.metadata["file_id"] == record.file_id
        assert document.metadata["file_path"] == str(record.absolute_path)
        assert document.metadata["relative_path"] == str(record.relative_path)
        assert document.metadata["file_name"] == record.filename
        assert document.metadata["file_extension"] == record.extension
        assert document.metadata["file_hash"] == record.sha256
        assert document.metadata["duplicate_group_id"] == record.duplicate_group_id
        assert document.metadata["parent_folder"] == str(record.parent_folder)
        assert document.metadata["inferred_category"] == record.inferred_category


def test_docx_preserves_section_and_table_source_locators(tmp_path):
    path = tmp_path / "structured-policy.docx"
    docx = DocxDocument()
    docx.add_heading("Возврат брака", level=1)
    docx.add_paragraph("Акт оформляется до передачи на склад.")
    table = docx.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Поле"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Ответственный"
    table.cell(1, 1).text = "Склад"
    docx.add_heading("Отгрузка", level=1)
    docx.add_paragraph("Водитель получает маршрутный лист.")
    docx.save(path)
    record = _record_for(path)

    result = extract_file(record)

    assert result.status == FileStatus.INDEXED
    assert [document.metadata["source_type"] for document in result.documents] == ["body", "table", "body"]
    first_body, table_doc, second_body = result.documents
    assert first_body.metadata["section_heading"] == "Возврат брака"
    assert first_body.metadata["source_locator"] == "section:1:возврат-брака:element:1"
    assert first_body.metadata["source_doc_id"].startswith(f"{record.file_id}:body:section:1")
    assert first_body.metadata["element_id"]
    assert first_body.metadata["element_hash"]
    assert "Возврат брака" in first_body.page_content
    assert "Акт оформляется" in first_body.page_content

    assert table_doc.metadata["section_heading"] == "Возврат брака"
    assert table_doc.metadata["table_index"] == 1
    assert table_doc.metadata["row_start"] == 1
    assert table_doc.metadata["row_end"] == 2
    assert table_doc.metadata["source_locator"] == "section:1:возврат-брака:element:2:table:1:rows:1-2"
    assert "Ответственный | Склад" in table_doc.page_content

    assert second_body.metadata["section_heading"] == "Отгрузка"
    assert second_body.metadata["source_locator"] == "section:2:отгрузка:element:3"
    assert "Водитель получает маршрутный лист." in second_body.page_content


def test_docx_detects_conservative_numbered_manual_heading(tmp_path):
    path = tmp_path / "manual-heading.docx"
    docx = DocxDocument()
    docx.add_paragraph("1. Область применения")
    docx.add_paragraph("Правило применяется ко всем складам.")
    docx.save(path)

    result = extract_file(_record_for(path))

    assert result.documents[0].metadata["section_heading"] == "1. Область применения"
    assert result.documents[0].metadata["element_ordinal"] == 1


def test_docx_without_text_is_no_text(tmp_path):
    path = tmp_path / "empty.docx"
    DocxDocument().save(path)
    record = _record_for(path)

    result = extract_file(record)

    assert result.status == FileStatus.NO_TEXT
    assert result.documents == []
    assert result.extraction_method == "python_docx"


def test_docx_embedded_images_are_ocrd_without_losing_body_text(tmp_path):
    embedded = tmp_path / "embedded.jpg"
    _make_image(embedded)
    path = tmp_path / "with-image.docx"
    docx = DocxDocument()
    docx.add_paragraph("Основной текст")
    docx.add_picture(str(embedded))
    docx.save(path)
    record = _record_for(path)

    result = extract_file(record, ocr_client=FakeOcrClient(), artifact_root=tmp_path / "artifacts")

    source_types = [document.metadata["source_type"] for document in result.documents]
    assert source_types == ["body", "embedded_image"]
    assert result.documents[1].metadata["image_index"] == 1
    assert result.documents[1].metadata["ocr_method"] == "fake_ocr"


def test_docx_embedded_image_ocr_failure_does_not_break_text_extraction(tmp_path):
    class FailingOcrClient:
        def extract_image_text(self, image_path: Path) -> OcrResult:
            raise RuntimeError("vision unavailable")

    embedded = tmp_path / "embedded.jpg"
    _make_image(embedded)
    path = tmp_path / "with-failing-image.docx"
    docx = DocxDocument()
    docx.add_paragraph("Текст остается доступен")
    docx.add_picture(str(embedded))
    docx.save(path)
    record = _record_for(path)

    result = extract_file(record, ocr_client=FailingOcrClient(), artifact_root=tmp_path / "artifacts")

    assert result.status == FileStatus.INDEXED
    assert [document.metadata["source_type"] for document in result.documents] == ["body"]
    assert "Текст остается доступен" in result.documents[0].page_content


def test_standalone_image_uses_ocr_client_and_cache(tmp_path):
    path = tmp_path / "scan.jpg"
    _make_image(path)
    record = _record_for(path)
    cache = OcrCache(tmp_path / "processed")
    first_client = FakeOcrClient(prefix="FIRST")

    first = extract_file(record, ocr_client=first_client, ocr_cache=cache)
    second_client = FakeOcrClient(prefix="SECOND")
    second = extract_file(record, ocr_client=second_client, ocr_cache=cache)

    assert first.status == FileStatus.INDEXED
    assert first.documents[0].metadata["source_type"] == "image"
    assert first.documents[0].page_content == "FIRST:scan.jpg"
    assert len(first_client.calls) == 1
    assert second.status == FileStatus.INDEXED
    assert second.documents[0].page_content == "FIRST:scan.jpg"
    assert second.documents[0].metadata["ocr_cached"] is True
    assert second_client.calls == []
    assert (tmp_path / "processed" / "ocr_cache.sqlite3").exists()


def test_ocr_cache_recipe_change_invalidates_cached_result(tmp_path):
    path = tmp_path / "scan.jpg"
    _make_image(path)
    record = _record_for(path)

    class ConfiguredFake(FakeOcrClient):
        def __init__(self, model: str) -> None:
            super().__init__(prefix=model)
            self.settings = type("Settings", (), {"vision_model": model, "ocr_task": "text_extraction"})()

    with OcrCache(tmp_path / "processed") as cache:
        first_client = ConfiguredFake("model-v1")
        second_client = ConfiguredFake("model-v2")
        first = extract_file(record, ocr_client=first_client, ocr_cache=cache)
        second = extract_file(record, ocr_client=second_client, ocr_cache=cache)

    assert first.documents[0].page_content.startswith("model-v1")
    assert second.documents[0].page_content.startswith("model-v2")
    assert len(first_client.calls) == 1
    assert len(second_client.calls) == 1
    assert first.documents[0].metadata["ocr_recipe_hash"] != second.documents[0].metadata["ocr_recipe_hash"]


def test_canned_ocr_response_is_rejected_before_indexing(tmp_path):
    path = tmp_path / "scan.jpg"
    _make_image(path)

    result = extract_file(_record_for(path), ocr_client=CannedDashScopeOcrClient())

    assert result.status == FileStatus.NO_TEXT
    assert result.documents == []
    assert "canned_no_text_response" in result.message


def test_ocr_routes_named_scheme_as_diagram(tmp_path):
    path = tmp_path / "СХЕМА ВОЗВРАТА.jpg"
    _make_image(path)

    result = extract_file(_record_for(path), ocr_client=FakeOcrClient(prefix="diagram text with enough context"))

    assert result.documents[0].metadata["layout_route"] == "diagram"


def test_ocr_cache_context_manager_closes_connection(tmp_path):
    with OcrCache(tmp_path / "processed") as cache:
        cache.write("scan", OcrResult(text="OCR text", method="fake"))
        cached = cache.read("scan")
        assert cached is not None
        assert cached.text == "OCR text"

    with pytest.raises(sqlite3.ProgrammingError):
        cache.read("scan")


def test_ocr_cache_finalizer_closes_unclosed_connection(monkeypatch, tmp_path):
    real_connect = sqlite3.connect
    opened: list[TrackingConnection] = []

    def tracking_connect(*args, **kwargs):
        connection = TrackingConnection(real_connect(*args, **kwargs))
        opened.append(connection)
        return connection

    monkeypatch.setattr(ocr_module.sqlite3, "connect", tracking_connect)
    cache = OcrCache(tmp_path / "processed")

    del cache

    assert opened
    assert all(connection.closed for connection in opened)


def test_standalone_image_without_ocr_text_is_no_text(tmp_path):
    path = tmp_path / "scan.png"
    _make_image(path)
    record = _record_for(path)

    result = extract_file(record, ocr_client=EmptyOcrClient())

    assert result.status == FileStatus.NO_TEXT
    assert result.documents == []
    assert result.extraction_method == "image_ocr"
    assert "OCR returned empty text for image" in result.message


def test_pdf_records_ocr_empty_pages_after_ocr_attempt(tmp_path):
    fitz = pytest.importorskip("fitz")
    path = tmp_path / "blank-scan.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(path)
    record = _record_for(path)

    result = extract_file(record, ocr_client=EmptyOcrClient(), artifact_root=tmp_path / "artifacts")

    assert result.status == FileStatus.NO_TEXT
    assert result.documents == []
    assert result.extraction_method == "pymupdf"
    assert "OCR returned empty text for page 1" in result.message


def test_pdf_extracts_native_text_and_ocrs_image_only_pages(tmp_path):
    fitz = pytest.importorskip("fitz")
    path = tmp_path / "scan.pdf"
    pdf = fitz.open()
    text_page = pdf.new_page()
    text_page.insert_text((72, 72), "Native PDF text")
    pdf.new_page()
    pdf.save(path)
    record = _record_for(path)

    result = extract_file(
        record,
        ocr_client=FakeOcrClient(),
        ocr_cache=OcrCache(tmp_path / "processed"),
        artifact_root=tmp_path / "artifacts",
    )

    assert result.status == FileStatus.INDEXED
    assert result.extraction_method == "pymupdf"
    assert [document.metadata["source_type"] for document in result.documents] == ["pdf_page", "pdf_page"]
    assert result.documents[0].metadata["page_number"] == 1
    assert "Native PDF text" in result.documents[0].page_content
    assert result.documents[1].metadata["page_number"] == 2
    assert result.documents[1].page_content.startswith("OCR:scan-page-2")


def test_xlsx_sheets_extract_rows_as_structured_text(tmp_path):
    path = tmp_path / "schedule.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "График"
    sheet.append(["Сотрудник", "Смена"])
    sheet.append(["Иванов", "Утро"])
    workbook.save(path)
    record = _record_for(path)

    result = extract_file(record)

    assert result.status == FileStatus.INDEXED
    assert result.extraction_method == "openpyxl"
    assert result.documents[0].metadata["source_type"] == "sheet"
    assert result.documents[0].metadata["sheet_name"] == "График"
    assert "Сотрудник | Смена" in result.documents[0].page_content
    assert "Иванов | Утро" in result.documents[0].page_content


def test_rtf_extracts_text(tmp_path):
    path = tmp_path / "note.rtf"
    path.write_text(r"{\rtf1\ansi Регламент склада}", encoding="utf-8")
    record = _record_for(path)

    result = extract_file(record)

    assert result.status == FileStatus.INDEXED
    assert result.extraction_method == "striprtf"
    assert "Регламент склада" in result.documents[0].page_content


def test_unsupported_extension_returns_unsupported(tmp_path):
    path = tmp_path / "notes.xyz"
    path.write_text("unsupported", encoding="utf-8")
    record = _record_for(path)

    result = extract_file(record)

    assert result.status == FileStatus.UNSUPPORTED
    assert result.documents == []
    assert "unsupported extension" in result.message
