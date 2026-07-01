from __future__ import annotations

import json
from collections import Counter
from pathlib import Path
from typing import Any, cast

from docx import Document as DocxDocument
from langchain_core.documents import Document
from openpyxl import Workbook
from PIL import Image

from imperial_rag.config import Settings
from imperial_rag.ingestion.manifest import FileStatus, IndexStatus, ManifestStore
from imperial_rag.ingestion.ocr import OcrResult
from imperial_rag.ingestion.pipeline import ingest_corpus


class DeterministicOcrClient:
    def __init__(self) -> None:
        self.calls: list[Path] = []

    def extract_image_text(self, image_path: Path) -> OcrResult:
        self.calls.append(image_path)
        return OcrResult(
            text="OCR_SENTINEL_SCAN_TEXT detected from pipeline image",
            method="deterministic_test_ocr",
        )


class FakeKeywordSearchIndex:
    last_settings: Settings | None = None
    last_documents: list[Any] | None = None

    def __init__(self, settings: Settings) -> None:
        FakeKeywordSearchIndex.last_settings = settings

    def replace_all(self, documents: list[Any]) -> None:
        FakeKeywordSearchIndex.last_documents = list(documents)


def test_real_pipeline_indexes_mixed_corpus_and_audits_failures(tmp_path: Path, monkeypatch) -> None:
    docs = tmp_path / "documents"
    docs.mkdir()
    _write_docx(docs / "policy.docx")
    _write_xlsx(docs / "schedule.xlsx")
    (docs / "note.rtf").write_text(r"{\rtf1\ansi RTF_SENTINEL warehouse rule}", encoding="utf-8")
    Image.new("RGB", (20, 20), "white").save(docs / "scan.png")
    (docs / "archive.rar").write_bytes(b"archive bytes are manifest-only in v1")
    (docs / "corrupted.docx").write_bytes(b"not a valid docx zip package")
    settings = Settings(workspace_root=tmp_path)
    ocr_client = DeterministicOcrClient()
    FakeKeywordSearchIndex.last_settings = None
    FakeKeywordSearchIndex.last_documents = None
    monkeypatch.setattr("imperial_rag.retrieval.elasticsearch.ElasticsearchKeywordIndex", FakeKeywordSearchIndex)

    summary = ingest_corpus(settings=settings, ocr_client=ocr_client, vector_store=None)

    assert summary.total_files == 6
    assert summary.indexed_files == 4
    assert summary.manifest_only_files == 1
    assert summary.failed_files == 1
    assert summary.no_text_files == 0
    assert summary.unsupported_files == 0
    assert summary.chunk_count == 5
    assert summary.keyword_indexed is True
    assert summary.vector_indexed is False
    assert len(ocr_client.calls) == 1

    manifest_store = ManifestStore(settings.manifest_db_path)
    records = {
        record.relative_path.as_posix(): record
        for record in manifest_store.list_records()
    }
    manifest_store.close()
    assert set(records) == {
        "policy.docx",
        "schedule.xlsx",
        "note.rtf",
        "scan.png",
        "archive.rar",
        "corrupted.docx",
    }

    expected_manifest = {
        "policy.docx": (FileStatus.INDEXED, "python_docx", 2),
        "schedule.xlsx": (FileStatus.INDEXED, "openpyxl", 1),
        "note.rtf": (FileStatus.INDEXED, "striprtf", 1),
        "scan.png": (FileStatus.INDEXED, "image_ocr", 1),
        "archive.rar": (FileStatus.MANIFEST_ONLY, None, 0),
        "corrupted.docx": (FileStatus.FAILED, "python_docx", 0),
    }
    for relative_path, (status, method, chunk_count) in expected_manifest.items():
        record = records[relative_path]
        assert record.status == status
        assert record.extraction_method == method
        assert record.chunk_count == chunk_count
        if status == FileStatus.INDEXED:
            assert record.keyword_index_status == IndexStatus.INDEXED
            assert record.vector_index_status == IndexStatus.SKIPPED
        else:
            assert record.keyword_index_status == IndexStatus.SKIPPED
            assert record.vector_index_status == IndexStatus.SKIPPED

    archive_error = records["archive.rar"].error_message
    assert archive_error is not None
    assert "archive files recorded but not extracted" in archive_error
    corrupted_error = records["corrupted.docx"].error_message
    assert corrupted_error
    assert any(fragment in corrupted_error.casefold() for fragment in ("package", "zip"))

    chunk_rows = _read_chunk_rows(settings.extraction_root / "chunks.jsonl")
    assert len(chunk_rows) == 5
    assert _relative_path_counts(chunk_rows) == {
        "policy.docx": 2,
        "schedule.xlsx": 1,
        "note.rtf": 1,
        "scan.png": 1,
    }
    for row in chunk_rows:
        metadata = row["metadata"]
        assert {
            "file_id",
            "relative_path",
            "source_type",
            "chunk_index",
            "chunk_id",
            "citation_id",
        } <= metadata.keys()

    chunk_text = "\n".join(row["page_content"] for row in chunk_rows)
    assert "DOCX_BODY_SENTINEL" in chunk_text
    assert "DOCX_TABLE_SENTINEL" in chunk_text
    assert "XLSX_SENTINEL" in chunk_text
    assert "RTF_SENTINEL" in chunk_text
    assert "OCR_SENTINEL_SCAN_TEXT" in chunk_text

    keyword_documents = FakeKeywordSearchIndex.last_documents
    assert FakeKeywordSearchIndex.last_settings is settings
    assert keyword_documents is not None
    keyword_documents = cast(list[Document], keyword_documents)
    assert len(keyword_documents) == 5
    assert Counter(document.metadata["relative_path"] for document in keyword_documents) == {
        "policy.docx": 2,
        "schedule.xlsx": 1,
        "note.rtf": 1,
        "scan.png": 1,
    }
    keyword_text = "\n".join(document.page_content for document in keyword_documents)
    assert "DOCX_BODY_SENTINEL" in keyword_text
    assert "DOCX_TABLE_SENTINEL" in keyword_text
    assert "XLSX_SENTINEL" in keyword_text
    assert "RTF_SENTINEL" in keyword_text
    assert "OCR_SENTINEL_SCAN_TEXT" in keyword_text


def _write_docx(path: Path) -> None:
    docx = DocxDocument()
    docx.add_paragraph("DOCX_BODY_SENTINEL return policy body text.")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "DOCX_TABLE_SENTINEL"
    table.cell(0, 1).text = "warehouse approval"
    docx.save(str(path))


def _write_xlsx(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Schedule"
    sheet.append(["Employee", "Shift"])
    sheet.append(["XLSX_SENTINEL", "Morning"])
    workbook.save(str(path))


def _read_chunk_rows(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines()]


def _relative_path_counts(rows: list[dict]) -> dict[str, int]:
    return dict(Counter(row["metadata"]["relative_path"] for row in rows))
